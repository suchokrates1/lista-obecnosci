from docx import Document
from docx.shared import Pt, Cm
import logging
import os

def ensure_template_exists(path: str) -> None:
    """Raise ``RuntimeError`` if template ``path`` does not exist."""
    if not os.path.exists(path):
        logger.error("Template file not found: %s", path)
        raise RuntimeError(f"Template file not found: {path}")

logger = logging.getLogger(__name__)

def generuj_liste_obecnosci(data, czas, obecni, trener, podpis_path, nazwa_zajec=None):
    template = "szablon.docx"
    ensure_template_exists(template)
    doc = Document(template)

    for para in doc.paragraphs:
        if "Lista obecności" in para.text and nazwa_zajec:
            para.text = f"Lista obecności – {nazwa_zajec}"
            continue
        if "Data zajęć:" in para.text and "Czas trwania zajęć:" in para.text:
            para.text = f"Data zajęć: {data}    Czas trwania zajęć: {czas}"

    if len(doc.tables) >= 2:
        tabela_uczestnicy = doc.tables[0]
        for i in range(1, len(tabela_uczestnicy.rows)):
            if i - 1 < len(obecni):
                cell = tabela_uczestnicy.cell(i, 0)
                cell.text = obecni[i - 1]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            else:
                for cell in tabela_uczestnicy.row_cells(i):
                    cell.text = ""

        tabela_trener = doc.tables[1]
        tabela_trener.cell(1, 0).text = trener
        if podpis_path and os.path.exists(podpis_path):
            try:
                run = tabela_trener.cell(1, 1).paragraphs[0].clear().add_run()
                run.add_picture(podpis_path, width=Cm(3.5))
            except Exception:
                logger.exception("Błąd przy podpisie")

    return doc

def generuj_raport_miesieczny(prowadzacy, zajecia, szablon_path, podpis_dir, miesiac, rok):
    ensure_template_exists(szablon_path)
    doc = Document(szablon_path)
    logger.debug("Generowanie raportu dla miesiąca: %s rok: %s", miesiac, rok)

    for para in doc.paragraphs:
        if "zleceniobiorca:" in para.text.lower():
            para.text = f"Zleceniobiorca: {prowadzacy.imie} {prowadzacy.nazwisko}"
        elif "zlecenia nr" in para.text.lower():
            para.text = f"Rozliczenie liczby godzin wykonywania usług do umowy zlecenia nr {prowadzacy.numer_umowy}"
        elif para.text.strip().lower().startswith("w "):
            para.text = f"w {miesiac:02d}.{rok}"

    dni_miesiaca = [d for d in zajecia if d.data.month == miesiac and d.data.year == rok]
    logger.debug("Liczba zajęć w miesiącu: %s", len(dni_miesiaca))

    suma = 0
    for tabela in doc.tables:
        for row in tabela.rows:
            dzien = None
            for col_idx, cell in enumerate(row.cells):
                for par in cell.paragraphs:
                    txt = par.text.strip()
                    logger.debug("Zawartość komórki: '%s'", txt)
                    txt_clean = txt.rstrip(".").strip().lstrip("0")
                    if txt_clean.isdigit():
                        dzien = int(txt_clean)
                        logger.debug("Sprawdzam dzień: %s", dzien)
            if dzien:
                laczny_czas = sum(
                    float(str(zaj.czas_trwania).replace(",", ".")) for zaj in dni_miesiaca if zaj.data.day == dzien
                )
                if laczny_czas > 0:
                    logger.debug("Łączny czas zajęć dla dnia %s: %s", dzien, laczny_czas)
                    godziny_txt = str(laczny_czas).replace(".0", "") + "h"
                    if len(row.cells) >= 3:
                        row.cells[1].text = godziny_txt
                        podpis_path = os.path.join(podpis_dir, prowadzacy.podpis_filename)
                        if os.path.exists(podpis_path):
                            try:
                                row.cells[2].paragraphs[0].clear().add_run().add_picture(podpis_path, width=Cm(2.5))
                            except Exception:
                                logger.exception("Błąd podpisu przy dniu %s", dzien)
                    suma += laczny_czas

    for row in doc.tables[-1].rows:
        zawiera_lacznie = any("łącznie" in cell.text.lower() for cell in row.cells)
        if zawiera_lacznie:
            for idx, cell in enumerate(row.cells):
                if "łącznie" in cell.text.lower():
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = str(suma).replace(".0", "") + "h"
                    if idx + 2 < len(row.cells):
                        podpis_path = os.path.join(podpis_dir, prowadzacy.podpis_filename)
                        if os.path.exists(podpis_path):
                            try:
                                row.cells[idx + 2].paragraphs[0].clear().add_run().add_picture(podpis_path, width=Cm(2.5))
                            except Exception:
                                logger.exception("Błąd podpisu w wierszu 'Łącznie'")

    for par in doc.paragraphs:
        if "(czytelny podpis zleceniobiorcy)" in par.text.lower():
            podpis_path = os.path.join(podpis_dir, prowadzacy.podpis_filename)
            if os.path.exists(podpis_path):
                try:
                    par.clear().add_run().add_picture(podpis_path, width=Cm(3.5))
                except Exception:
                    logger.exception("Błąd podpisu na dole")

    return doc

