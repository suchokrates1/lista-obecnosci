import os
import io
from PIL import Image
import pytest
import json
import subprocess
from app import create_app
from datetime import datetime, timedelta
from model import db, Uzytkownik, Prowadzacy, Zajecia, Uczestnik, PasswordResetToken, Setting
from docx import Document
import utils
from werkzeug.security import generate_password_hash
import re


@pytest.fixture
def app(tmp_path):
    os.environ["SECRET_KEY"] = "testsecret"
    os.environ["DATABASE_URL"] = "sqlite:///" + str(tmp_path / "test.db")
    os.environ["MAX_SIGNATURE_SIZE"] = "10"
    os.environ["SMTP_HOST"] = "smtp"
    os.environ["SMTP_PORT"] = "25"
    os.environ["EMAIL_LOGIN"] = "user"
    os.environ["EMAIL_PASSWORD"] = "pass"
    app = create_app()
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        yield app
        db.session.remove()
        db.drop_all()


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture(scope="session", autouse=True)
def ensure_jsdom():
    if not os.path.exists("node_modules/jsdom"):
        subprocess.run(["npm", "install", "jsdom@22"], check=True)


def test_start_fails_without_mail_vars(tmp_path, monkeypatch):
    monkeypatch.delenv("SMTP_HOST", raising=False)
    monkeypatch.delenv("SMTP_PORT", raising=False)
    monkeypatch.delenv("EMAIL_LOGIN", raising=False)
    monkeypatch.delenv("EMAIL_PASSWORD", raising=False)
    os.environ["SECRET_KEY"] = "x"
    os.environ["DATABASE_URL"] = "sqlite:///" + str(tmp_path / "db.db")
    os.environ["MAX_SIGNATURE_SIZE"] = "10"
    with pytest.raises(RuntimeError):
        create_app()


def test_routes_accessible(client):
    assert client.get("/login").status_code == 200
    assert client.get("/register").status_code == 200
    resp = client.get("/")
    assert resp.status_code == 302
    assert "/login" in resp.headers["Location"]


def test_register_missing_fields(client, app):
    resp = client.post("/register", data={}, follow_redirects=False)
    assert resp.status_code == 302
    with app.app_context():
        assert Uzytkownik.query.count() == 0


def test_register_invalid_email(client, app):
    data = {
        "imie": "A",
        "nazwisko": "B",
        "numer_umowy": "1",
        "uczestnik": ["X"],
        "login": "invalid",
        "haslo": "pass",
    }
    resp = client.post("/register", data=data, follow_redirects=False)
    assert resp.status_code == 302
    with app.app_context():
        assert Uzytkownik.query.count() == 0


def test_register_bad_signature(client, app):
    data = {
        "imie": "A",
        "nazwisko": "B",
        "numer_umowy": "1",
        "uczestnik": ["X"],
        "login": "a@example.com",
        "haslo": "pass",
        "podpis": (io.BytesIO(b"data"), "sig.txt", "text/plain"),
    }
    resp = client.post(
        "/register",
        data=data,
        content_type="multipart/form-data",
        follow_redirects=False,
    )
    assert resp.status_code == 302
    with app.app_context():
        assert Uzytkownik.query.count() == 0


def test_register_too_large_signature(client, app):
    big = io.BytesIO(b"0123456789ABCDEF")
    data = {
        "imie": "A",
        "nazwisko": "B",
        "numer_umowy": "1",
        "uczestnik": ["X"],
        "login": "b@example.com",
        "haslo": "pass",
        "podpis": (big, "sig.png", "image/png"),
    }
    resp = client.post(
        "/register",
        data=data,
        content_type="multipart/form-data",
        follow_redirects=False,
    )
    assert resp.status_code == 302
    with app.app_context():
        assert Uzytkownik.query.count() == 0


def test_successful_register(client, app, monkeypatch):
    monkeypatch.setattr("routes.auth.send_plain_email", lambda *a, **k: None)
    monkeypatch.setattr(utils, "SIGNATURE_MAX_SIZE", 1000)
    buf = io.BytesIO()
    Image.new("RGB", (1, 1), (255, 0, 0)).save(buf, format="PNG")
    buf.seek(0)
    data = {
        "imie": "A",
        "nazwisko": "B",
        "numer_umowy": "1",
        "uczestnik": ["X"],
        "login": "ok@example.com",
        "haslo": "pass",
        "podpis": (buf, "sig.png", "image/png"),
    }
    resp = client.post(
        "/register",
        data=data,
        content_type="multipart/form-data",
        follow_redirects=False,
    )
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/login")
    with app.app_context():
        assert Uzytkownik.query.count() == 1


def test_register_with_new_participant_fields(client, app, monkeypatch):
    monkeypatch.setattr("routes.auth.send_plain_email", lambda *a, **k: None)
    from werkzeug.datastructures import MultiDict

    data = MultiDict(
        [
            ("imie", "A"),
            ("nazwisko", "B"),
            ("numer_umowy", "1"),
            ("uczestnik", "X"),
            ("uczestnik", "Y\nZ"),
            ("login", "new@example.com"),
            ("haslo", "pass"),
        ]
    )
    resp = client.post(
        "/register",
        data=data,
        content_type="multipart/form-data",
        follow_redirects=False,
    )
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/login")
    with app.app_context():
        prow = Prowadzacy.query.first()
        names = [u.imie_nazwisko for u in prow.uczestnicy]
        assert names == ["X", "Y", "Z"]


def test_login_success(client, app):
    with app.app_context():
        user = Uzytkownik(
            login="adm@example.com",
            haslo_hash=generate_password_hash("secret"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.commit()
    resp = client.post(
        "/login",
        data={"login": "adm@example.com", "hasło": "secret"},
        follow_redirects=False,
    )
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/admin")


def test_login_remember_sets_cookie(client, app):
    with app.app_context():
        user = Uzytkownik(
            login="perm@example.com",
            haslo_hash=generate_password_hash("secret"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.commit()

    resp = client.post(
        "/login",
        data={"login": "perm@example.com", "hasło": "secret", "remember": "1"},
        follow_redirects=False,
    )
    assert resp.status_code == 302
    cookie = client._cookies.get(("localhost", "/", "remember_token"))
    assert cookie is not None
    assert cookie.expires is not None


def test_login_failure(client, app):
    with app.app_context():
        user = Uzytkownik(
            login="adm2@example.com",
            haslo_hash=generate_password_hash("secret"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.commit()
    resp = client.post(
        "/login",
        data={"login": "adm2@example.com", "hasło": "bad"},
        follow_redirects=False,
    )
    assert resp.status_code == 200
    assert b"Nieprawid" in resp.data


def test_admin_index_page(client, app):
    """Loading the attendance page as admin should work."""
    with app.app_context():
        prow = Prowadzacy(imie="A", nazwisko="B")
        db.session.add(prow)
        admin = Uzytkownik(
            login="idx@example.com",
            haslo_hash=generate_password_hash("pass"),
            role="admin",
            approved=True,
        )
        db.session.add(admin)
        db.session.commit()

    client.post(
        "/login",
        data={"login": "idx@example.com", "hasło": "pass"},
        follow_redirects=False,
    )
    resp = client.get("/")
    assert resp.status_code == 200


def test_panel_requires_login(client):
    resp = client.get("/panel")
    assert resp.status_code == 302
    assert "/login" in resp.headers["Location"]


def test_panel_raport_requires_login(client):
    resp = client.get("/panel/raport")
    assert resp.status_code == 302
    assert "/login" in resp.headers["Location"]


def test_admin_raport_bad_params(client, app):
    """Invalid month/year should return 400 for the admin route."""
    with app.app_context():
        prow = Prowadzacy(imie="A", nazwisko="B")
        db.session.add(prow)
        user = Uzytkownik(
            login="adm@example.com",
            haslo_hash=generate_password_hash("secret"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.commit()
        prow_id = prow.id
        zaj = Zajecia(
            prowadzacy_id=prow_id, data=datetime(2023, 1, 1), czas_trwania=1.0
        )
        db.session.add(zaj)
        db.session.commit()
    client.post("/login", data={"login": "adm@example.com", "hasło": "secret"})
    resp = client.get(f"/raport/{prow_id}?miesiac=13&rok=2023")
    assert resp.status_code == 400


def test_panel_raport_bad_params(client, app):
    """Invalid month/year should redirect with an error for the trainer route."""
    with app.app_context():
        prow = Prowadzacy(imie="T", nazwisko="X")
        db.session.add(prow)
        user = Uzytkownik(
            login="tr@example.com",
            haslo_hash=generate_password_hash("secret"),
            role="prowadzacy",
            approved=True,
            prowadzacy=prow,
        )
        db.session.add(user)
        db.session.commit()
        zaj = Zajecia(
            prowadzacy_id=prow.id, data=datetime(2023, 1, 1), czas_trwania=1.0
        )
        db.session.add(zaj)
        db.session.commit()
    client.post("/login", data={"login": "tr@example.com", "hasło": "secret"})
    resp = client.get("/panel/raport?rok=1999")
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/panel")


def _create_trainer(app):
    """Create a trainer account with one session and return the login."""
    with app.app_context():
        prow = Prowadzacy(
            imie="T", nazwisko="T", numer_umowy="1", podpis_filename="sig.png"
        )
        db.session.add(prow)
        db.session.flush()
        user = Uzytkownik(
            login="t@example.com",
            haslo_hash=generate_password_hash("pass"),
            role="prowadzacy",
            approved=True,
            prowadzacy_id=prow.id,
        )
        db.session.add(user)
        zaj = Zajecia(
            prowadzacy_id=prow.id,
            data=datetime(2023, 5, 1),
            czas_trwania=1.0,
        )
        db.session.add(zaj)
        db.session.commit()
        return user.login


@pytest.fixture
def trainer(client, app):
    """Create a trainer with one participant and log in."""
    login_val = _create_trainer(app)
    with app.app_context():
        prow = Prowadzacy.query.first()
        prow.uczestnicy.append(Uczestnik(imie_nazwisko="Osoba"))
        db.session.commit()
    client.post(
        "/login",
        data={"login": login_val, "hasło": "pass"},
        follow_redirects=False,
    )
    return login_val


def test_panel_raport_access_control(client, app):
    with app.app_context():
        user = Uzytkownik(
            login="a@example.com",
            haslo_hash=generate_password_hash("x"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.commit()
    client.post(
        "/login", data={"login": "a@example.com", "hasło": "x"}, follow_redirects=False
    )
    resp = client.get("/panel/raport")
    assert resp.status_code == 403


def test_panel_raport_download(client, app, monkeypatch):
    login_val = _create_trainer(app)

    def dummy_report(*_a, **_k):
        doc = Document()
        doc.add_paragraph("test")
        return doc

    monkeypatch.setattr("routes.panel.generuj_raport_miesieczny", dummy_report)

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.get("/panel/raport")
    assert resp.status_code == 200
    assert (
        resp.mimetype
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in resp.headers.get("Content-Disposition", "")
    assert "raport_" in resp.headers.get("Content-Disposition", "")
    assert len(resp.data) > 0


def test_panel_raport_email_sending(client, app, monkeypatch):
    login_val = _create_trainer(app)

    def dummy_report(*_a, **_k):
        doc = Document()
        doc.add_paragraph("test")
        return doc

    sent = {}

    def fake_email(buf, data, typ=None):
        sent["called"] = True

    monkeypatch.setattr("routes.panel.generuj_raport_miesieczny", dummy_report)
    monkeypatch.setattr("routes.panel.email_do_koordynatora", fake_email)

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.get("/panel/raport?wyslij=1")
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/panel")
    assert sent.get("called")


def test_wyslij_zajecie_requires_login(client):
    resp = client.get("/wyslij_zajecie/1")
    assert resp.status_code == 302
    assert "/login" in resp.headers["Location"]


def test_wyslij_zajecie_success(client, app, monkeypatch):
    login_val = _create_trainer(app)

    with app.app_context():
        zaj_id = Zajecia.query.first().id

    called = {}

    def fake_send(z):
        called["sent"] = True
        z.wyslano = True
        db.session.commit()
        return True

    monkeypatch.setattr("routes.panel.send_attendance_list", fake_send)

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.get(f"/wyslij_zajecie/{zaj_id}")
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/panel")
    assert called.get("sent")
    with app.app_context():
        assert db.session.get(Zajecia, zaj_id).wyslano is True

    # Check that the session row shows the sent indicator
    resp = client.get("/panel")
    assert resp.status_code == 200
    data = resp.data.decode()
    assert "bi bi-check-lg text-success" in data


def test_wyslij_zajecie_admin_requires_login(client):
    resp = client.get("/wyslij_zajecie_admin/1")
    assert resp.status_code == 302
    assert "/login" in resp.headers["Location"]


def test_wyslij_zajecie_admin_success(client, app, monkeypatch):
    _ = _create_trainer(app)

    with app.app_context():
        admin = Uzytkownik(
            login="admin@example.com",
            haslo_hash=generate_password_hash("adm"),
            role="admin",
            approved=True,
        )
        db.session.add(admin)
        db.session.commit()
        zaj_id = Zajecia.query.first().id

    called = {}

    def fake_send(z):
        called["sent"] = True
        z.wyslano = True
        db.session.commit()
        return True

    monkeypatch.setattr("routes.admin.send_attendance_list", fake_send)

    client.post(
        "/login",
        data={"login": "admin@example.com", "hasło": "adm"},
        follow_redirects=False,
    )
    resp = client.get(f"/wyslij_zajecie_admin/{zaj_id}")
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/admin")
    assert called.get("sent")
    with app.app_context():
        assert db.session.get(Zajecia, zaj_id).wyslano is True

    resp = client.get("/admin")
    assert resp.status_code == 200
    data = resp.data.decode()
    assert "bi bi-check-lg text-success" in data


def test_wyslij_zajecie_admin_requires_admin(client, app):
    login_val = _create_trainer(app)

    with app.app_context():
        zaj_id = Zajecia.query.first().id

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.get(f"/wyslij_zajecie_admin/{zaj_id}")
    assert resp.status_code == 403


def test_admin_dashboard_filter(client, app):
    with app.app_context():
        p1 = Prowadzacy(imie="A", nazwisko="A")
        p2 = Prowadzacy(imie="B", nazwisko="B")
        db.session.add_all([p1, p2])
        db.session.flush()
        z1 = Zajecia(prowadzacy_id=p1.id, data=datetime(2023, 1, 1), czas_trwania=1.0)
        z2 = Zajecia(prowadzacy_id=p2.id, data=datetime(2023, 1, 2), czas_trwania=2.0)
        db.session.add_all([z1, z2])
        admin = Uzytkownik(
            login="adminf@example.com",
            haslo_hash=generate_password_hash("adm"),
            role="admin",
            approved=True,
        )
        db.session.add(admin)
        db.session.commit()
        pid = p1.id

    client.post(
        "/login",
        data={"login": "adminf@example.com", "hasło": "adm"},
        follow_redirects=False,
    )
    resp = client.get(f"/admin?p_id={pid}")
    assert resp.status_code == 200
    data = resp.data.decode()
    assert "2023-01-01" in data
    assert "2023-01-02" not in data


def test_update_default_time(client, app):
    login_val = _create_trainer(app)
    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.post(
        "/panel/profil",
        data={
            "imie": "T",
            "nazwisko": "T",
            "numer_umowy": "1",
            "domyslny_czas": "2,5",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302
    with app.app_context():
        prow = Prowadzacy.query.first()
        assert prow.domyslny_czas == 2.5


def test_usun_uczestnika_requires_login(client, app):
    with app.app_context():
        prow = Prowadzacy(imie="A")
        db.session.add(prow)
        db.session.flush()
        u = Uczestnik(imie_nazwisko="X", prowadzacy_id=prow.id)
        db.session.add(u)
        db.session.commit()
        uid = u.id

    resp = client.post(f"/usun_uczestnika/{uid}")
    assert resp.status_code == 302
    assert "/login" in resp.headers["Location"]


def test_usun_uczestnika_forbidden(client, app):
    login_val = _create_trainer(app)
    with app.app_context():
        prow2 = Prowadzacy(imie="B")
        db.session.add(prow2)
        db.session.flush()
        user2 = Uzytkownik(
            login="t2@example.com",
            haslo_hash=generate_password_hash("x"),
            role="prowadzacy",
            approved=True,
            prowadzacy_id=prow2.id,
        )
        db.session.add(user2)
        u = Uczestnik(imie_nazwisko="Z", prowadzacy_id=prow2.id)
        db.session.add(u)
        db.session.commit()
        uid = u.id

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.post(f"/usun_uczestnika/{uid}")
    assert resp.status_code == 403


def test_trainer_delete_participant(client, app):
    login_val = _create_trainer(app)
    with app.app_context():
        prow = Prowadzacy.query.first()
        u = Uczestnik(imie_nazwisko="P", prowadzacy_id=prow.id)
        db.session.add(u)
        db.session.commit()
        uid = u.id

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.post(f"/usun_uczestnika/{uid}", follow_redirects=False)
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith("/panel")
    with app.app_context():
        assert db.session.get(Uczestnik, uid) is None


def test_add_and_rename_participant(client, app):
    """Trainer can add and rename a participant using panel routes."""

    login_val = _create_trainer(app)
    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )

    resp = client.post(
        "/panel/dodaj_uczestnika",
        data={"new_participant": "Nowy"},
        follow_redirects=False,
    )
    assert resp.status_code == 302
    with app.app_context():
        prow = Prowadzacy.query.first()
        u = Uczestnik.query.filter_by(
            prowadzacy_id=prow.id, imie_nazwisko="Nowy"
        ).first()
        assert u is not None
        uid = u.id

    resp = client.post(
        f"/panel/zmien_uczestnika/{uid}",
        data={"new_name": "Zmieniony"},
        follow_redirects=False,
    )
    assert resp.status_code == 302
    with app.app_context():
        assert db.session.get(Uczestnik, uid).imie_nazwisko == "Zmieniony"


def test_admin_add_participant(client, app):
    with app.app_context():
        prow = Prowadzacy(imie="A", nazwisko="B")
        db.session.add(prow)
        admin = Uzytkownik(
            login="admadd@example.com",
            haslo_hash=generate_password_hash("x"),
            role="admin",
            approved=True,
        )
        db.session.add(admin)
        db.session.commit()
        pid = prow.id

    client.post(
        "/login",
        data={"login": "admadd@example.com", "hasło": "x"},
        follow_redirects=False,
    )
    resp = client.post(
        f"/admin/trainer/{pid}/add_participant",
        data={"new_participant": "P"},
        follow_redirects=False,
    )
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith(f"/admin/trainer/{pid}")
    with app.app_context():
        u = Uczestnik.query.filter_by(prowadzacy_id=pid, imie_nazwisko="P").first()
        assert u is not None
        uid = u.id


def test_admin_delete_participant(client, app):
    with app.app_context():
        prow = Prowadzacy(imie="C", nazwisko="D")
        db.session.add(prow)
        part = Uczestnik(imie_nazwisko="Q", prowadzacy=prow)
        admin = Uzytkownik(
            login="admdel@example.com",
            haslo_hash=generate_password_hash("x"),
            role="admin",
            approved=True,
        )
        db.session.add_all([admin, part])
        db.session.commit()
        pid = prow.id
        uid = part.id

    client.post(
        "/login",
        data={"login": "admdel@example.com", "hasło": "x"},
        follow_redirects=False,
    )
    resp = client.post(f"/admin/participant/{uid}/delete", follow_redirects=False)
    assert resp.status_code == 302
    assert resp.headers["Location"].endswith(f"/admin/trainer/{pid}")
    with app.app_context():
        assert db.session.get(Uczestnik, uid) is None


def test_reset_request_purges_expired_token(client, app, monkeypatch):
    with app.app_context():
        user = Uzytkownik(
            login="purge@example.com",
            haslo_hash=generate_password_hash("x"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.flush()
        uid = user.id
        t = PasswordResetToken(
            user_id=user.id,
            token="old",
            expires_at=datetime.utcnow() - timedelta(hours=2),
        )
        db.session.add(t)
        db.session.commit()

    monkeypatch.setattr("routes.auth.send_plain_email", lambda *a, **k: None)
    resp = client.post("/reset-request", data={"login": "purge@example.com"})
    assert resp.status_code == 302
    with app.app_context():
        tokens = PasswordResetToken.query.filter_by(user_id=uid).all()
        assert len(tokens) == 1
        assert tokens[0].token != "old"


def test_reset_with_token_purges_expired(client, app):
    with app.app_context():
        user = Uzytkownik(
            login="tok@example.com",
            haslo_hash=generate_password_hash("x"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.flush()
        expired = PasswordResetToken(
            user_id=user.id,
            token="old",
            expires_at=datetime.utcnow() - timedelta(hours=1),
        )
        valid = PasswordResetToken(
            user_id=user.id,
            token="good",
            expires_at=datetime.utcnow() + timedelta(hours=1),
        )
        db.session.add_all([expired, valid])
        db.session.commit()
        token_value = valid.token

    resp = client.get(f"/reset/{token_value}")
    assert resp.status_code == 200
    with app.app_context():
        assert PasswordResetToken.query.filter_by(token="old").first() is None
        assert PasswordResetToken.query.filter_by(token=token_value).first() is not None


def test_panel_summary_table_links(client, app):
    login_val = _create_trainer(app)
    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )
    resp = client.get("/panel")
    assert resp.status_code == 200
    data = resp.data.decode()
    assert "Raporty miesięczne" in data
    assert "/panel/raport?rok=2023&miesiac=5" in data
    assert "/panel/raport?rok=2023&miesiac=5&wyslij=1" in data



def test_admin_statystyki_requires_admin(client, app):
    login_val = _create_trainer(app)
    with app.app_context():
        trainer = Prowadzacy.query.first()
        user = Uzytkownik(
            login="t2@example.com",
            haslo_hash=generate_password_hash("x"),
            role="prowadzacy",
            approved=True,
            prowadzacy_id=trainer.id,
        )
        db.session.add(user)
        db.session.commit()
        tid = trainer.id

    client.post(
        "/login", data={"login": "t2@example.com", "hasło": "x"}, follow_redirects=False
    )
    resp = client.get(f"/admin/statystyki/{tid}")
    assert resp.status_code == 403


def test_admin_statystyki_data(client, app):
    _ = _create_trainer(app)
    with app.app_context():
        trainer = Prowadzacy.query.first()
        u1 = Uczestnik(imie_nazwisko="A", prowadzacy_id=trainer.id)
        u2 = Uczestnik(imie_nazwisko="B", prowadzacy_id=trainer.id)
        admin = Uzytkownik(
            login="admstat@example.com",
            haslo_hash=generate_password_hash("a"),
            role="admin",
            approved=True,
        )
        db.session.add_all([u1, u2, admin])
        zaj = Zajecia.query.first()
        zaj.obecni.append(u1)
        db.session.commit()
        tid = trainer.id

    client.post(
        "/login",
        data={"login": "admstat@example.com", "hasło": "a"},
        follow_redirects=False,
    )
    resp = client.get(f"/admin/statystyki/{tid}")
    assert resp.status_code == 200
    data = resp.data.decode()
    assert "A" in data and "1/1" in data
    assert "B" in data and "0/1" in data


def _create_many_sessions(app, user_login: str, count: int) -> None:
    """Add ``count`` dummy sessions for the trainer with ``user_login``."""
    with app.app_context():
        user = Uzytkownik.query.filter_by(login=user_login).first()
        assert user is not None
        for i in range(count):
            db.session.add(
                Zajecia(
                    prowadzacy_id=user.prowadzacy_id,
                    data=datetime(2023, 1, i + 1),
                    czas_trwania=1.0,
                )
            )
        db.session.commit()


def test_panel_pagination(client, app):
    login_val = _create_trainer(app)
    _create_many_sessions(app, login_val, 12)

    client.post(
        "/login", data={"login": login_val, "hasło": "pass"}, follow_redirects=False
    )

    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert "2023-01-12" in html
    assert "2023-01-02" not in html

    resp2 = client.get("/panel?page=2")
    assert resp2.status_code == 200
    html2 = resp2.data.decode()
    assert "2023-01-02" in html2
    assert "2023-01-12" not in html2


def test_admin_pagination(client, app):
    login_val = _create_trainer(app)
    _create_many_sessions(app, login_val, 12)
    with app.app_context():
        admin = Uzytkownik(
            login="pageadm@example.com",
            haslo_hash=generate_password_hash("a"),
            role="admin",
            approved=True,
        )
        db.session.add(admin)
        db.session.commit()

    client.post(
        "/login",
        data={"login": "pageadm@example.com", "hasło": "a"},
        follow_redirects=False,
    )

    resp = client.get("/admin")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert "2023-01-12" in html
    assert "2023-01-02" not in html

    resp2 = client.get("/admin?page=2")
    assert resp2.status_code == 200
    html2 = resp2.data.decode()
    assert "2023-01-02" in html2
    assert "2023-01-12" not in html2


def test_panel_progress_and_edit_forms(client, trainer):
    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert "progress-bar" in html
    assert "Frekwencja %" in html
    assert "<table" in html

    resp2 = client.get("/panel?edit=1")
    assert resp2.status_code == 200
    html2 = resp2.data.decode()
    assert "name=\"new_participant\"" in html2
    assert "name=\"new_name\"" in html2

def test_panel_profile_edit_mode(client, trainer):
    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert "name=\"imie\"" not in html
    assert "name=\"nazwisko\"" not in html
    assert "name=\"domyslny_czas\"" not in html

    resp2 = client.get("/panel?edit_profile=1")
    assert resp2.status_code == 200
    html2 = resp2.data.decode()
    assert "name=\"imie\"" in html2
    assert "name=\"nazwisko\"" in html2
    assert "name=\"domyslny_czas\"" in html2
    assert "name=\"podpis\"" in html2


def test_panel_edit_profile_shows_form_fields(client, trainer):
    resp = client.get("/panel?edit_profile=1")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert "name=\"imie\"" in html
    assert "name=\"nazwisko\"" in html
    assert "name=\"domyslny_czas\"" in html


def test_panel_profile_post_updates_trainer(client, app, trainer):
    data = {
        "imie": "Nowe",
        "nazwisko": "Nazwisko",
        "numer_umowy": "99",
        "domyslny_czas": "3",
    }
    resp = client.post("/panel/profil", data=data, follow_redirects=False)
    assert resp.status_code == 302
    with app.app_context():
        prow = Prowadzacy.query.first()
        assert prow.imie == "Nowe"
        assert prow.nazwisko == "Nazwisko"
        assert prow.numer_umowy == "99"
        assert prow.domyslny_czas == 3.0


def test_panel_edit_history_inputs(client, trainer):
    resp = client.get("/panel?edit=1")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'name="data"' in html
    assert 'name="czas"' in html


def test_admin_edit_mode_inputs(client, app):
    with app.app_context():
        p = Prowadzacy(imie="A", nazwisko="B", numer_umowy="1")
        db.session.add(p)
        db.session.flush()
        z = Zajecia(prowadzacy_id=p.id, data=datetime(2023, 1, 1), czas_trwania=1.0)
        admin = Uzytkownik(login="admtest@example.com", haslo_hash=generate_password_hash("x"), role="admin", approved=True)
        db.session.add_all([z, admin])
        db.session.commit()
    client.post("/login", data={"login": "admtest@example.com", "hasło": "x"})
    resp = client.get("/admin?edit=1")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'name="imie"' in html
    assert 'name="data"' in html


def test_admin_stats_edit_mode_inputs(client, app):
    with app.app_context():
        p = Prowadzacy(imie="S", nazwisko="T")
        db.session.add(p)
        db.session.flush()
        u = Uczestnik(imie_nazwisko="P", prowadzacy_id=p.id)
        z = Zajecia(prowadzacy_id=p.id, data=datetime(2023, 1, 1), czas_trwania=1.0)
        z.obecni.append(u)
        admin = Uzytkownik(login="admstat@example.com", haslo_hash=generate_password_hash("s"), role="admin", approved=True)
        db.session.add_all([u, z, admin])
        db.session.commit()
        pid = p.id
    client.post("/login", data={"login": "admstat@example.com", "hasło": "s"})
    resp = client.get(f"/admin/statystyki/{pid}?edit=1")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'name="present_' in html


def _login_admin(client, app):
    with app.app_context():
        user = Uzytkownik(
            login="admw@example.com",
            haslo_hash=generate_password_hash("pass"),
            role="admin",
            approved=True,
        )
        db.session.add(user)
        db.session.commit()
        utils.load_db_settings(app)
    client.post("/login", data={"login": "admw@example.com", "hasło": "pass"})


def _parse_settings_preview(html: str) -> dict:
    proc = subprocess.run(
        ["node", "tests/run_theme.js"],
        input=html.encode(),
        capture_output=True,
        check=True,
    )
    return json.loads(proc.stdout.decode())


def test_save_column_widths(client, app):
    _login_admin(client, app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_admin_trainers_id": "10",
            "width_admin_trainers_name": "30",
            "width_admin_trainers_signature": "20",
            "width_admin_trainers_participants": "20",
            "width_admin_trainers_action": "20",
            "width_panel_history_sent": "20",
            "width_panel_history_date": "20",
            "width_panel_history_duration": "20",
            "width_panel_history_participants": "20",
            "width_panel_history_action": "20",
            "width_panel_profile_data_first": "20",
            "width_panel_profile_data_last": "20",
            "width_panel_profile_data_contract": "20",
            "width_panel_profile_data_default": "20",
            "width_panel_profile_data_signature": "20",
            "width_panel_participants_name": "20",
            "width_panel_participants_percent": "40",
            "width_panel_participants_present": "40",
            "width_panel_monthly_reports_year": "25",
            "width_panel_monthly_reports_month": "25",
            "width_panel_monthly_reports_hours": "25",
            "width_panel_monthly_reports_action": "25",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302
    with app.app_context():
        setting = Setting.query.get("table_admin_trainers_widths")
        assert setting is not None
        assert "id=10.0" in setting.value
        assert "name=30.0" in setting.value
        hist = Setting.query.get("table_panel_history_widths")
        assert hist is not None
        parts = {}
        for item in hist.value.split(','):
            if not item:
                continue
            key, val = item.split('=')
            parts[key] = float(val)
        assert abs(sum(parts.values()) - 100.0) < 0.1
        assert parts.get("date") == 20.0
        profile = Setting.query.get("table_panel_profile_data_widths")
        assert profile is not None
        assert "first=20.0" in profile.value
        participants = Setting.query.get("table_panel_participants_widths")
        assert participants is not None
        parts_p = {}
        for item in participants.value.split(','):
            if not item:
                continue
            key, val = item.split('=')
            parts_p[key] = float(val)
        assert abs(sum(parts_p.values()) - 100.0) < 0.1
        assert parts_p.get("percent") == 40.0
        monthly = Setting.query.get("table_panel_monthly_reports_widths")
        assert monthly is not None
        parts_m = {}
        for item in monthly.value.split(','):
            if not item:
                continue
            key, val = item.split('=')
            parts_m[key] = float(val)
        assert abs(sum(parts_m.values()) - 100.0) < 0.1
        assert parts_m.get("month") == 25.0


def test_save_column_widths_invalid_sum(client, app):
    _login_admin(client, app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_admin_trainers_id": "50",
            "width_admin_trainers_name": "30",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 200
    assert b"Suma" in resp.data
    with app.app_context():
        assert Setting.query.get("table_admin_trainers_widths") is None


def test_admin_page_applies_column_widths(client, app):
    _login_admin(client, app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_admin_trainers_id": "25",
            "width_admin_trainers_name": "25",
            "width_admin_trainers_signature": "25",
            "width_admin_trainers_participants": "15",
            "width_admin_trainers_action": "10",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302
    utils.load_db_settings(app)
    resp = client.get("/admin")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'col-admin-trainers-id" style="width: 25.0%' in html


def test_panel_profile_widths_render(client, app):
    _login_admin(client, app)
    trainer_login = _create_trainer(app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_panel_profile_data_first": "10",
            "width_panel_profile_data_last": "20",
            "width_panel_profile_data_contract": "20",
            "width_panel_profile_data_default": "30",
            "width_panel_profile_data_signature": "20",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302
    utils.load_db_settings(app)
    client.get("/logout")
    client.post("/login", data={"login": trainer_login, "hasło": "pass"})
    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'col-panel-profile-data-first" style="width: 10.0%' in html


def test_panel_participants_widths_render(client, app):
    _login_admin(client, app)
    trainer_login = _create_trainer(app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_panel_participants_name": "30",
            "width_panel_participants_percent": "30",
            "width_panel_participants_present": "40",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302
    utils.load_db_settings(app)
    client.get("/logout")
    client.post("/login", data={"login": trainer_login, "hasło": "pass"})
    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'col-panel-participants-name" style="width: 30.0%' in html


def test_panel_monthly_reports_widths_render(client, app):
    _login_admin(client, app)
    trainer_login = _create_trainer(app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_panel_monthly_reports_year": "40",
            "width_panel_monthly_reports_month": "20",
            "width_panel_monthly_reports_hours": "20",
            "width_panel_monthly_reports_action": "20",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302
    utils.load_db_settings(app)
    client.get("/logout")
    client.post("/login", data={"login": trainer_login, "hasło": "pass"})
    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    assert 'col-panel-monthly-reports-year" style="width: 40.0%' in html


def test_admin_settings_rejects_bad_widths(client, app):
    with app.app_context():
        db.session.add(Setting(key="table_admin_trainers_widths", value="id=20,name=80"))
        db.session.commit()

    _login_admin(client, app)
    resp = client.post(
        "/admin/settings",
        data={"width_admin_trainers_id": "50", "width_admin_trainers_name": "30"},
        follow_redirects=False,
    )
    assert resp.status_code == 200
    assert b"Suma" in resp.data
    with app.app_context():
        setting = Setting.query.get("table_admin_trainers_widths")
        assert setting.value == "id=20,name=80"


def test_settings_preview_shows_saved_widths(client, app):
    _login_admin(client, app)
    resp = client.post(
        "/admin/settings",
        data={
            "width_admin_trainers_id": "25",
            "width_admin_trainers_name": "75",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 302

    resp = client.get("/admin/settings")
    assert resp.status_code == 200
    result = _parse_settings_preview(resp.data.decode())
    assert result["idWidth"].startswith("25")
    assert result["nameWidth"].startswith("75")
    assert result["warn"] == ""


def test_theme_totals_warning_zero_once_valid():
    html = (
        "<table id='admin-trainers'><col id='admin-trainers-id'><col id='admin-trainers-name'></table>"
        "<span class='total-warning' data-table='admin_trainers'></span>"
        "<input data-table='admin_trainers' data-column='id' value='60'>"
        "<input data-table='admin_trainers' data-column='name' value='30'>"
    )
    invalid = _parse_settings_preview(html)
    assert invalid["warn"] != ""
    html_valid = html.replace("value='30'", "value='40'")
    valid = _parse_settings_preview(html_valid)
    assert valid["warn"] == ""


def test_panel_tables_have_table_bordered(client, trainer):
    resp = client.get("/panel")
    assert resp.status_code == 200
    html = resp.data.decode()
    for tid in ("panel-participants", "panel-monthly-reports", "panel-history"):
        m = re.search(r'<table[^>]*id="%s"[^>]*>' % tid, html)
        assert m and "table-bordered" in m.group(0)


def test_admin_history_table_bordered(client, app):
    _create_trainer(app)
    _login_admin(client, app)
    resp = client.get("/admin")
    assert resp.status_code == 200
    html = resp.data.decode()
    m = re.search(r'<table[^>]*id="admin-history"[^>]*>', html)
    assert m and "table-bordered" in m.group(0)


def test_admin_stats_table_bordered(client, app):
    _create_trainer(app)
    with app.app_context():
        tid = Prowadzacy.query.first().id
    _login_admin(client, app)
    resp = client.get(f"/admin/statystyki/{tid}")
    assert resp.status_code == 200
    html = resp.data.decode()
    m = re.search(r'<table[^>]*id="admin-stats"[^>]*>', html)
    assert m and "table-bordered" in m.group(0)
