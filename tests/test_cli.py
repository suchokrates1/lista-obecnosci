import os
from datetime import datetime

import pytest
from docx import Document

from app import create_app
from model import db, Prowadzacy, Zajecia


@pytest.fixture
def app(tmp_path):
    os.environ["SECRET_KEY"] = "test"
    os.environ["DATABASE_URL"] = "sqlite:///" + str(tmp_path / "db.db")
    os.environ["MAX_SIGNATURE_SIZE"] = "10"
    os.environ["SMTP_HOST"] = "smtp"
    os.environ["SMTP_PORT"] = "25"
    os.environ["EMAIL_LOGIN"] = "user"
    os.environ["EMAIL_PASSWORD"] = "pass"
    application = create_app()
    application.config["WTF_CSRF_ENABLED"] = False
    with application.app_context():
        db.create_all()
        yield application
        db.session.remove()
        db.drop_all()


def _setup_data(app):
    with app.app_context():
        p1 = Prowadzacy(imie="A", nazwisko="B", podpis_filename="sig.png")
        p2 = Prowadzacy(imie="C", nazwisko="D", podpis_filename="sig.png")
        db.session.add_all([p1, p2])
        db.session.flush()
        z1 = Zajecia(prowadzacy_id=p1.id, data=datetime(2025, 5, 1), czas_trwania=1.0)
        z2 = Zajecia(prowadzacy_id=p2.id, data=datetime(2025, 6, 1), czas_trwania=1.0)
        db.session.add_all([z1, z2])
        db.session.commit()
        return p1.id, p2.id


def test_generate_reports_file(app, monkeypatch, tmp_path):
    p1_id, p2_id = _setup_data(app)

    def dummy_report(*_a, **_k):
        doc = Document()
        doc.add_paragraph("x")
        return doc

    monkeypatch.setattr("app.generuj_raport_miesieczny", dummy_report)
    runner = app.test_cli_runner()
    monkeypatch.chdir(tmp_path)
    result = runner.invoke(args=["generate-reports", "--month", "5", "--year", "2025"])
    assert result.exit_code == 0
    expected = tmp_path / "reports" / f"raport_{p1_id}_5_2025.docx"
    assert expected.exists()
    not_expected = tmp_path / "reports" / f"raport_{p2_id}_5_2025.docx"
    assert not not_expected.exists()


def test_generate_reports_email(app, monkeypatch, tmp_path):
    _setup_data(app)

    def dummy_report(*_a, **_k):
        doc = Document()
        doc.add_paragraph("x")
        return doc

    sent = {}

    def fake_email(buf, data, typ=None, course=None):
        sent["called"] = True

    monkeypatch.setattr("app.generuj_raport_miesieczny", dummy_report)
    monkeypatch.setattr("app.email_do_koordynatora", fake_email)
    runner = app.test_cli_runner()
    monkeypatch.chdir(tmp_path)
    result = runner.invoke(
        args=["generate-reports", "--month", "5", "--year", "2025", "--email"]
    )
    assert result.exit_code == 0
    assert sent.get("called")
