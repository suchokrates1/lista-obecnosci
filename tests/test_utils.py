import io
from werkzeug.datastructures import FileStorage
from PIL import Image
from datetime import datetime, timedelta
import os
from pathlib import Path
import smtplib
from docx import Document
from email.message import EmailMessage

import utils
import pytest
from utils import is_valid_email
from model import db, Uzytkownik, PasswordResetToken, Prowadzacy, Zajecia
from werkzeug.security import generate_password_hash
from app import create_app


@pytest.fixture
def app(tmp_path):
    os.environ['SECRET_KEY'] = 'testsecret'
    os.environ['DATABASE_URL'] = 'sqlite:///' + str(tmp_path / 'test.db')
    os.environ['MAX_SIGNATURE_SIZE'] = '10'
    os.environ['SMTP_HOST'] = 'smtp'
    os.environ['SMTP_PORT'] = '25'
    os.environ['EMAIL_LOGIN'] = 'user'
    os.environ['EMAIL_PASSWORD'] = 'pass'
    Path("szablon.docx").touch()
    Path("rejestr.docx").touch()
    application = create_app()
    application.config['WTF_CSRF_ENABLED'] = False
    try:
        with application.app_context():
            db.create_all()
            yield application
            db.session.remove()
            db.drop_all()
    finally:
        Path("szablon.docx").unlink(missing_ok=True)
        Path("rejestr.docx").unlink(missing_ok=True)


@pytest.fixture
def client(app):
    return app.test_client()


def test_is_valid_email():
    assert is_valid_email('user@example.com')
    assert not is_valid_email('invalid')
    assert not is_valid_email('a@')


def test_validate_signature_ok(monkeypatch):
    monkeypatch.setattr(utils, 'SIGNATURE_MAX_SIZE', 10)
    fs = FileStorage(io.BytesIO(b'x'), filename='sig.png', content_type='image/png')
    name, error = utils.validate_signature(fs)
    assert name == 'sig.png'
    assert error is None


def test_validate_signature_bad_extension(monkeypatch):
    monkeypatch.setattr(utils, 'SIGNATURE_MAX_SIZE', 10)
    fs = FileStorage(io.BytesIO(b'x'), filename='sig.txt', content_type='text/plain')
    name, error = utils.validate_signature(fs)
    assert name is None
    assert error


def test_validate_signature_too_big(monkeypatch):
    monkeypatch.setattr(utils, 'SIGNATURE_MAX_SIZE', 1)
    fs = FileStorage(io.BytesIO(b'ab'), filename='sig.png', content_type='image/png')
    name, error = utils.validate_signature(fs)
    assert name is None
    assert error


def test_process_signature(monkeypatch, tmp_path):
    monkeypatch.setattr(utils, 'REMOVE_SIGNATURE_BG', False)
    monkeypatch.chdir(tmp_path)
    (tmp_path / 'static').mkdir()
    img = Image.new('RGB', (1, 1), (255, 0, 0))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    filename = utils.process_signature(buf)
    saved = tmp_path / 'static' / filename
    assert saved.exists()
    out = Image.open(saved)
    assert out.format == 'PNG'

def test_validate_signature_none():
    name, error = utils.validate_signature(None)
    assert name is None
    assert error is None


class FailingStream(io.BytesIO):
    def seek(self, *a, **k):
        raise IOError('fail')


def test_validate_signature_exception():
    fs = FileStorage(FailingStream(b'x'), filename='sig.png', content_type='image/png')
    with pytest.raises(utils.SignatureValidationError):
        utils.validate_signature(fs)


def test_send_plain_email_queue(monkeypatch):
    called = {}

    def fake_send(msg):
        called['to'] = msg['To']

    monkeypatch.setattr(utils, '_send_message', fake_send)
    utils.send_plain_email(
        'x@example.com',
        'SUBJ',
        'BODY',
        's',
        'b',
        queue=True,
    )
    utils.shutdown_email_worker()
    assert called.get('to') == 'x@example.com'
    assert utils._worker is None or not utils._worker.is_alive()

def test_shutdown_email_worker_thread(monkeypatch):
    utils.shutdown_email_worker()
    monkeypatch.setattr(utils, "_send_message", lambda msg: None)
    utils.send_plain_email(
        "y@example.com",
        "SUBJ",
        "BODY",
        "s",
        "b",
        queue=True,
    )
    assert utils._worker is not None and utils._worker.is_alive()
    utils.shutdown_email_worker()
    assert utils._worker is None or not utils._worker.is_alive()


def test_safe_format_missing(monkeypatch):
    assert utils.safe_format("x {foo} {bar}", foo="y") == "x y "


def test_send_plain_email_extra_placeholder(monkeypatch):
    called = {}

    def fake_send(msg):
        called["subject"] = msg["Subject"]
        called["body"] = msg.get_content()

    monkeypatch.setattr(utils, "_send_message", fake_send)
    os.environ["EXTRA_SUBJ"] = "A {name} {missing}"
    os.environ["EXTRA_BODY"] = "B {missing}"
    utils.send_plain_email(
        "z@example.com",
        "EXTRA_SUBJ",
        "EXTRA_BODY",
        "s",
        "b",
        name="Bob",
    )
    assert called["subject"] == "A Bob "
    assert called["body"].startswith("B ")


def test_email_do_koordynatora_extra_placeholder(monkeypatch, app):
    with app.app_context():
        os.environ["EMAIL_RECIPIENT"] = "coord@example.com"
        os.environ["EMAIL_LOGIN"] = "user@example.com"
        os.environ["EMAIL_LIST_SUBJECT"] = "Sub {date} {missing}"
        os.environ["EMAIL_LIST_BODY"] = "Body {missing}"

        called = {}

        def fake_send(msg):
            called["subject"] = msg["Subject"]
            called["body"] = msg.get_content()

        monkeypatch.setattr(utils, "_send_message", fake_send)

        buf = io.BytesIO(b"x")
        utils.email_do_koordynatora(buf, "2025-01-01")

        assert called["subject"].startswith("Sub 2025-01-01")
        assert called["body"].startswith("Body ")



def test_purge_expired_tokens(app):
    with app.app_context():
        user = Uzytkownik(
            login='exp@example.com',
            haslo_hash=generate_password_hash('x'),
            role='admin',
            approved=True,
        )
        db.session.add(user)
        db.session.flush()
        past = datetime.utcnow() - timedelta(hours=2)
        future = datetime.utcnow() + timedelta(hours=1)
        t1 = PasswordResetToken(user_id=user.id, token='old', expires_at=past)
        t2 = PasswordResetToken(user_id=user.id, token='new', expires_at=future)
        db.session.add_all([t1, t2])
        db.session.commit()

        utils.purge_expired_tokens()

        assert PasswordResetToken.query.filter_by(token='old').first() is None
        assert PasswordResetToken.query.filter_by(token='new').first() is not None


def _create_simple_session():
    prow = Prowadzacy(imie="T", nazwisko="T", numer_umowy="1", nazwa_zajec="Z", podpis_filename="sig.png")
    db.session.add(prow)
    db.session.flush()
    zaj = Zajecia(prowadzacy_id=prow.id, data=datetime(2023, 5, 1), czas_trwania=1.0)
    db.session.add(zaj)
    db.session.commit()
    return zaj


def test_send_attendance_list_success(app, monkeypatch):
    with app.app_context():
        zaj = _create_simple_session()

        def dummy_doc(*_a, **_k):
            doc = Document()
            doc.add_paragraph("x")
            return doc

        called = {}

        def fake_email(buf, data, typ="lista", course=None, queue=False, trainer=None):
            called["sent"] = True
            called["trainer"] = trainer

        monkeypatch.setattr(utils, "generuj_liste_obecnosci", dummy_doc)
        monkeypatch.setattr(utils, "email_do_koordynatora", fake_email)

        assert utils.send_attendance_list(zaj)
        assert called.get("sent")
        assert called.get("trainer") is zaj.prowadzacy
        assert zaj.wyslano is True


def test_send_attendance_list_failure(app, monkeypatch):
    with app.app_context():
        zaj = _create_simple_session()

        def dummy_doc(*_a, **_k):
            doc = Document()
            doc.add_paragraph("x")
            return doc

        def fail_email(buf, data, typ="lista", course=None, queue=False, trainer=None):
            raise smtplib.SMTPException("x")

        monkeypatch.setattr(utils, "generuj_liste_obecnosci", dummy_doc)
        monkeypatch.setattr(utils, "email_do_koordynatora", fail_email)

        assert utils.send_attendance_list(zaj) is False
        assert zaj.wyslano is False


def test_email_do_koordynatora_trainer_name(monkeypatch, app):
    with app.app_context():
        prow = Prowadzacy(imie="Jan", nazwisko="Kowalski", numer_umowy="1", nazwa_zajec="Z", podpis_filename="sig.png")
        db.session.add(prow)
        db.session.commit()

        os.environ["EMAIL_RECIPIENT"] = "coord@example.com"
        os.environ["EMAIL_LOGIN"] = "user@example.com"
        os.environ["EMAIL_USE_TRAINER_NAME"] = "1"

        called = {}

        def fake_send(msg):
            called["from"] = msg["From"]

        monkeypatch.setattr(utils, "_send_message", fake_send)

        buf = io.BytesIO(b"x")
        utils.email_do_koordynatora(buf, "2025-01-01", trainer=prow)

        assert called["from"].startswith("Jan Kowalski <")


def test_email_do_koordynatora_default_name(monkeypatch, app):
    with app.app_context():
        prow = Prowadzacy(imie="Jan", nazwisko="Kowalski", numer_umowy="1", nazwa_zajec="Z", podpis_filename="sig.png")
        db.session.add(prow)
        db.session.commit()

        os.environ["EMAIL_RECIPIENT"] = "coord@example.com"
        os.environ["EMAIL_LOGIN"] = "user@example.com"
        os.environ["EMAIL_SENDER_NAME"] = "Vest"
        os.environ["EMAIL_USE_TRAINER_NAME"] = "0"

        called = {}

        def fake_send(msg):
            called["from"] = msg["From"]

        monkeypatch.setattr(utils, "_send_message", fake_send)

        buf = io.BytesIO(b"x")
        utils.email_do_koordynatora(buf, "2025-01-01", trainer=prow)

        assert called["from"].startswith("Vest <")


def test_month_name_filter_registered(app):
    """The month_name Jinja filter should return Polish month names."""
    filt = app.jinja_env.filters.get("month_name")
    assert filt is not None
    assert filt(1) == "styczeń"
    assert filt("12") == "grudzień"


def test_parse_registration_form_old_field(app):
    from werkzeug.datastructures import MultiDict

    form = MultiDict(
        {
            "imie": "A",
            "nazwisko": "B",
            "numer_umowy": "1",
            "nazwa_zajec": "Z",
            "lista_uczestnikow": "X\nY",
            "login": "u@example.com",
            "haslo": "pass",
        }
    )
    data, error = utils.parse_registration_form(form, MultiDict())
    assert error is None
    assert data["uczestnicy"] == ["X", "Y"]


def test_parse_registration_form_new_field(app):
    from werkzeug.datastructures import MultiDict

    form = MultiDict(
        [
            ("imie", "A"),
            ("nazwisko", "B"),
            ("numer_umowy", "1"),
            ("nazwa_zajec", "Z"),
            ("uczestnik", "X"),
            ("uczestnik", "Y\nZ"),
            ("login", "u2@example.com"),
            ("haslo", "pass"),
        ]
    )
    data, error = utils.parse_registration_form(form, MultiDict())
    assert error is None
    assert data["uczestnicy"] == ["X", "Y", "Z"]


def test_attach_cid_images_no_escape(monkeypatch, tmp_path):
    msg = EmailMessage()
    msg.set_content("<p>x</p>", subtype="html")
    (tmp_path / "static").mkdir()
    img = tmp_path / "static" / "ok.png"
    img.write_bytes(b"x")
    monkeypatch.chdir(tmp_path)
    utils._attach_cid_images(msg, '<img src="cid:../ok.png">')
    assert len(list(msg.iter_attachments())) == 0
